from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


# ── Region detection ──────────────────────────────────────────────────────────


def detect_region(source: str) -> str:
    """
    Detect document region from job source tag.
    Returns: 'in' | 'gb' | 'de' | 'ae' | 'us'
    """
    source_lower = source.lower()
    if "/in" in source_lower or "india" in source_lower:
        return "in"
    if "/gb" in source_lower or "uk" in source_lower or "britain" in source_lower:
        return "gb"
    if "/de" in source_lower or "germany" in source_lower or "deutsch" in source_lower:
        return "de"
    if (
        "/ae" in source_lower
        or "uae" in source_lower
        or "dubai" in source_lower
        or "gulf" in source_lower
    ):
        return "ae"
    return "us"  # default


def get_page_config(region: str) -> dict:
    """
    Return page size and margin config for the region.
    US Letter for US, A4 for everything else.
    """
    if region == "us":
        return {
            "width": Inches(8.5),
            "height": Inches(11),
            "top": Inches(0.75),
            "bottom": Inches(0.75),
            "left": Inches(1.0),
            "right": Inches(1.0),
            "label": "US Letter",
        }
    else:
        # A4 for India, UK, Germany, UAE, Europe
        return {
            "width": Cm(21),
            "height": Cm(29.7),
            "top": Cm(1.5),
            "bottom": Cm(1.5),
            "left": Cm(2.0),
            "right": Cm(2.0),
            "label": "A4",
        }


# ── Core converter ────────────────────────────────────────────────────────────


def markdown_to_docx(md_text: str, output_path: str, source: str = ""):
    """
    Convert ATS resume markdown to a clean formatted .docx file.
    Auto-detects region from source for correct page size.
    Optimized for single-page / compact ATS resume output.
    """
    region = detect_region(source)
    page_cfg = get_page_config(region)

    doc = Document()

    # ── Page size and margins ─────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width = page_cfg["width"]
        section.page_height = page_cfg["height"]
        section.top_margin = page_cfg["top"]
        section.bottom_margin = page_cfg["bottom"]
        section.left_margin = page_cfg["left"]
        section.right_margin = page_cfg["right"]

    # ── Default font — compact for resume ─────────────────────────────────────
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10)  # 10pt keeps resume compact
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(2)
    normal_style.paragraph_format.line_spacing = Pt(13)

    # ── Process lines ─────────────────────────────────────────────────────────
    lines = md_text.split("\n")

    for line in lines:
        line = line.rstrip()

        if line.startswith("# "):
            # Candidate name — large, centered
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith("## "):
            # Section header with blue underline
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[3:].upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            _add_bottom_border(p)

        elif line.startswith("### "):
            # Role / project title
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        elif line.startswith("- "):
            # Bullet — compact indent
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.2)
            _add_inline_bold(p, line[2:])
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"

        elif (
            line.strip().startswith("**")
            and line.strip().endswith("**")
            and line.strip().count("**") == 2
        ):
            # Standalone bold line (e.g. skill category)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line.strip().strip("*"))
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        elif line == "" or line == "---":
            # Minimal spacing for blank lines
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)

        else:
            if line.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                _add_inline_bold(p, line)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

    doc.save(output_path)
    return output_path


# ── Helpers ───────────────────────────────────────────────────────────────────


def _add_bottom_border(paragraph):
    """Add a blue bottom border to a paragraph (section divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_inline_bold(paragraph, text):
    """Handle **bold** text inline within a paragraph."""
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for idx, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            run.bold = idx % 2 == 1
